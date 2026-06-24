from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path("E:/WorkingSpace/研究相关")
OUT = BASE / "0624盛誉周报.docx"


def set_run_font(run, name="宋体", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("盛誉周报")
    set_run_font(run, "黑体", 16, True)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("2026.06.18 - 2026.06.24")
    set_run_font(run, "宋体", 11)


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_run_font(run, "黑体", 13, True)


def add_paragraph(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(22)
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    set_run_font(run, "宋体", 11)


def build_report():
    doc = Document()
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(11)
    add_title(doc)

    sections = [
        (
            "周四-周五：整理训练入口和数据读取问题",
            [
                "这两天先把前面做 S-DeCI 时比较乱的测试入口重新整理了一遍。之前每次测 ABIDE、MDD、Mātai、Taowu 都要在 run_cv 或一堆参数里手动改，容易漏参数，所以我单独写了 test_abide_best_config.py、test_mdd_best_config.py、test_matai_best_config.py 和 test_taowu_best_config.py。每个脚本里都直接放了当前认为比较合适的一组默认参数，后面在 IDE 里直接运行就能测，不用每次重新拼命令。",
                "整理脚本时还顺手检查了数据集读取。这里发现一个很具体的问题：data_factory_CV.py 里 Taowu 原来映射到 Neurocon_Dataset，这样即使传入 data=Taowu，实际读取逻辑也不是 Taowu 自己的数据类。我把这个映射改成了 Taowu_Dataset，并用 1 fold、1 epoch 的方式快速跑了一遍，确认 Taowu 能读到 40 个样本，Mātai 能读到 60 个样本，数据形状分别是 [batch, 239, 116] 和 [batch, 200, 116]。",
                "我还重新检查了 K-fold 的划分逻辑。现在用的是 StratifiedKFold，shuffle=True，random_state 跟 seed 绑定，所以是从全部样本里随机分层划分训练集和测试集，不是按文件夹顺序切。这个检查主要是因为之前有些结果太整齐，我怀疑可能是数据划分或者标签传递有问题。",
                "另外，ABIDE300 样本数很少的问题也查清楚了。原因不是加载错了，而是现在时间序列读取逻辑改成了严格筛选 seq_len 及以上的样本。seq_len=300 时，ABIDE 里很多样本长度不足 300 会被跳过，最后只剩 138 个左右。之前我试过短序列补零，但后面觉得补零会改变时间序列频谱和相关矩阵，所以又改回了短序列直接过滤、长序列截断。这个逻辑对 ALFF/fALFF 和相关矩阵都更干净一点。",
            ],
        ),
        (
            "周末：把模块一改成更偏生理信号的输入",
            [
                "模块一这周主要是重新考虑输入特征。之前 DeCI/Cycle 的思路是把时间序列分解成周期和高频部分，但从可视化看，后面 z_global 经常学不到稳定类别差异，训练集能拟合，测试集分不开。我怀疑一开始输入给后面模块的特征就混了太多站点差异、扫描噪声或个体噪声。",
                "所以我去查了静息态 fMRI 里常用的 ALFF/fALFF 思路。ALFF 是看低频段振幅，fALFF 是低频振幅占整个频段振幅的比例，常用频段大概是 0.01-0.08 Hz。这种特征比直接把原始时间序列丢给模型更接近生理活动指标，也更符合我论文里“先减少站点偏移和噪声，再做因果图学习”的思路。",
                "具体操作上，我把模块一的 feature mode 增加到 alff 路线，里面会根据 TR、low Hz、high Hz 算低频振幅相关特征，同时保留一定时间统计信息。现在脚本里默认 module1_feature_mode=alff，module1_tr=2.0，module1_alff_low_hz=0.01，module1_alff_high_hz=0.08。这样后面模块拿到的不是完全原始的时间序列，而是先经过低频生理特征筛选的表示。",
                "我也开始比较不同数据集上的 ALFF 表示，看它相比原始信号是不是更稳定。现在还不能说它已经明显提高准确率，但至少从设计上更贴近 fMRI 的常用特征，不再是完全让网络自己从原始序列里学所有东西。",
            ],
        ),
        (
            "周一：重新改模块二的因果学习目标",
            [
                "模块二之前的问题比较明显：因果矩阵变化很小，有时候看起来更像一个被分类损失轻微调整的邻接矩阵，不像是真正学到了时间序列里的因果依赖。我先后参考了 NOTEARS、DAGMA 和 Differentiable DAG Sampling 的源码，发现如果只在静态特征上做图学习，很容易和下游分类纠缠在一起。",
                "因此这周我把模块二的目标往时间序列预测式 SEM 上改。也就是不再只用模块一输出的静态特征去重构自己，而是把模块一处理后的时间序列作为输入，用图结构去预测下一时刻或后续时刻的 ROI 信号。这样图的训练目标来自时间动态本身，而不是只来自分类是否正确。",
                "具体实现上，脚本里增加了 causal_learning_target=temporal_sem，temporal_lag_order=2，lambda_temporal_pred、lambda_temporal_sparse、lambda_temporal_smooth 分别控制预测误差、稀疏性和平滑性。DAGMA 的部分也做了 warmup 和 barrier 调度，避免一开始无环约束太强导致图完全学不动。",
                "我还把因果图可视化改成每个 fold 训练结束后保存一次，而不是固定若干 step 保存一次。这样更适合看一个 fold 最后到底学成什么样。同时保留 graph_mass、direction、dag_raw、delta_abs 等诊断值，方便判断图是不是过密、有没有方向性、有没有真正变化。",
            ],
        ),
        (
            "周二：继续调模块三和模块四，并加可视化",
            [
                "模块三还是沿着 HGCN 做，但我这周做了两个比较具体的修改。一个是 readout 从普通 node_stats 增加了 network_stats，按 AAL116 的脑网络分组去读出特征。这个改法的来源是我觉得单纯对 116 个 ROI 做全局平均会把很多脑网络层面的差异抹掉，而脑区网络级统计可能更适合分类。",
                "另一个是继续检查模块三得到的 z_global。之前只看训练集不够，因为训练集能分开不代表测试集能分开。所以我把最后一个 epoch 的训练集和测试集一起做 t-SNE，可视化时训练集和测试集用不同形状，标签用颜色区分，prototype 也画在图里。这样能直接看到测试样本是不是围着正确原型分布。",
                "模块四这周主要改成每类多原型。这个思路来自我之前看的 Prototypical Representation Learning for Multi-Site Domain Generalization in Schizophrenia Diagnosis，那篇文章里不是只学一个类别中心，而是通过原型约束让不同域和不同类的特征更稳定。我把每类原型数改成可调，比如 hpec_prototypes_per_class=5，并加入了 PCL、PAL、HSIC、类内原型正交和类间 margin 这些项。",
                "不过从目前图上看，原型仍然容易挤在一起，尤其测试集 z_global 和原型之间没有形成很干净的对应关系。所以我又调过 lambda_hpec_pcl、lambda_hpec_pal、lambda_hpec_hsic、lambda_hpec_intra_orthogonal 和 lambda_hpec_inter_margin，但提升不稳定。这个结果说明模块四不是简单多加几个 loss 就能解决，前面模块三输出的特征本身也要更有类别信息才行。",
            ],
        ),
        (
            "周三：几个数据集上的实验现象",
            [
                "这周跑实验时我主要看三件事：训练准确率、测试准确率、t-SNE 里测试集能不能分开。现在比较一致的现象是训练集经常能拟合得比较好，但测试集 z_global 分不开，说明过拟合还是比较明显。",
                "ABIDE300 上问题最明显。因为筛选 seq_len=300 后样本只剩一百多个，模型稍微复杂一点就很容易记住训练集，测试准确率有时甚至只有三十多。这个结果让我基本排除了“只是调参不够”的解释，因为最优 epoch 很低，而且继续训练测试集还会下降，说明模型学到的不是能泛化的疾病特征。",
                "MDD 上因为是多站点数据，我原本希望模块一的 ALFF/fALFF 和 site_zscore 能减少站点偏移，但现在结果也没有稳定到理想水平。可视化里训练集能更明显分开，测试集仍然混在一起，说明站点差异和标签差异还没有被很好地区分开。",
                "Mātai 和 Taowu 我主要先做了入口和流程验证。Mātai 用 1 fold、1 epoch 快速跑通时，最终验证 acc 约 0.83、macro F1 约 0.81；Taowu 同样快速跑通时 acc 约 0.625、macro F1 约 0.56。但这两个只是 smoke test，不是正式性能结论，后面还要完整跑 5 fold 和多次 iteration。",
                "总的来说，现在的模型能跑通、能输出中间量、能做消融，但分类能力还没有达到我希望的程度。尤其是 z_global 没有稳定区分标签，这说明问题更可能在模块一到模块三的表征学习链路，而不只是模块四的原型损失。",
            ],
        ),
        (
            "目前判断和下周要做的事",
            [
                "我现在的判断是：S-DeCI 的四个模块思路还可以保留，但需要让每个模块的训练目标更明确。模块一应该先把 fMRI 里和生理活动相关的低频信息提出来；模块二应该主要靠时间序列预测学图，而不是靠分类 loss 间接调图；模块三要让因果图或相关图真正影响 z_global，而不是最后仍然像普通分类器；模块四则要在 z_global 有区分度的前提下再做原型约束。",
                "下周我准备先不继续盲目调一堆参数，而是做更细的消融：只开模块一看 ALFF/fALFF 的分布；开模块一和二看因果图是否随训练明显变化；开模块一二三看 z_global 是否已经有类别趋势；最后再打开模块四看原型有没有把这种趋势增强。如果前面 z_global 本身就没有标签结构，后面再调 HPEC 意义不大。",
                "另外我还会继续看 ALFF/fALFF 和多站点偏移相关论文，重点不是再加一个新模块，而是把模块一做得更符合静息态 fMRI 的特征提取习惯。现在模型效果不好我会如实记录，但至少这周已经把训练入口、数据读取、模块可视化和几个关键结构问题都排查清楚了，后面可以更有针对性地改。",
            ],
        ),
    ]

    for title, paragraphs in sections:
        add_heading(doc, title)
        for text in paragraphs:
            add_paragraph(doc, text)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
